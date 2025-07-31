from langchain.chat_models import init_chat_model
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langgraph.graph import START, StateGraph
from typing_extensions import List, TypedDict
from langchain_core.prompts import PromptTemplate
from langchain_community.document_loaders import PyMuPDFLoader
import docx
from langchain_core.documents import Document
from PIL import Image
import pytesseract
import pandas as pd
import sqlite3
import os


# Loaders
def load_pdf(file_path):
    loader = PyMuPDFLoader(file_path)
    return loader.load()

def load_docx(file_path):
    doc = docx.Document(file_path)
    full_text = "\n".join([para.text for para in doc.paragraphs])
    return [Document(page_content=full_text, metadata={"source": file_path})]

def load_txt(file_path):
    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read()
    return [Document(page_content=text, metadata={"source": file_path})]

def load_image(file_path):
    text = pytesseract.image_to_string(Image.open(file_path))
    return [Document(page_content=text, metadata={"source": file_path})]

def load_csv(file_path):
    df = pd.read_csv(file_path)
    text = df.to_csv(index=False)
    return [Document(page_content=text, metadata={"source": file_path})]

def load_sqlite(file_path, query="SELECT name FROM sqlite_master WHERE type='table';"):
    conn = sqlite3.connect(file_path)
    cursor = conn.cursor()
    documents = []
    cursor.execute(query)
    tables = [row[0] for row in cursor.fetchall()]
    for table in tables:
        df = pd.read_sql_query(f"SELECT * FROM {table}", conn)
        text = df.to_csv(index=False)
        documents.append(Document(page_content=text, metadata={"source": f"{file_path}:{table}"}))
    conn.close()
    return documents

def load_xlsx(file_path):
    df = pd.read_excel(file_path)
    text = df.to_csv(index=False)
    return [Document(page_content=text, metadata={"source": file_path})]

def load_json(file_path):
    with open(file_path, "r", encoding="utf-8") as f:
        data = f.read()
    return [Document(page_content=data, metadata={"source": file_path})]

def load_md(file_path):
    with open(file_path, "r", encoding="utf-8") as f:
        markdown_text = f.read()
    return [Document(page_content=markdown_text, metadata={"source": file_path})]


def load_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return load_pdf(file_path)
    elif ext == ".docx":
        return load_docx(file_path)
    elif ext == ".txt":
        return load_txt(file_path)
    elif ext in [".jpg", ".jpeg", ".png"]:
        return load_image(file_path)
    elif ext == ".csv":
        return load_csv(file_path)
    elif ext in [".db", ".sqlite"]:
        return load_sqlite(file_path)
    elif ext == ".xlsx":
        return load_xlsx(file_path)
    elif ext == ".json":
        return load_json(file_path)
    elif ext == ".md":
        return load_md(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")




llm = init_chat_model("gemini-2.5-flash", model_provider="google_genai")

embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")


vector_store = Chroma(
    collection_name="rag_collection",
    embedding_function=embeddings,
)

files = ["file1.pdf", "file2.docx", "image1.png", "data.csv", "database.db"]
docs = []

for file in files:
    doc = load_file(file)
    docs.extend(doc)

text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=1000, 
    chunk_overlap=200,
    separators=["\n\n", "\n", " ", ""]
    )

all_splits = text_splitter.split_documents(docs)

vectors = vector_store.add_documents(documents=all_splits)


prompt_template = """Use the following pieces of context to answer the question.
If you don't know the answer based on the context that has been provided then, just say that you don't know, don't try to make up an answer.
You have to be concise with your answers so that its reasonable and understandable to the user.

{context}

Question: {question}

Helpful Answer:
"""

rag_prompt = PromptTemplate.from_template(prompt_template)


class State(TypedDict):
    question: str
    context: List[Document]
    answer: str


def retrieve(state: State):
    retrieved_docs = vectors.similarity_search(state["question"])
    return {"context": retrieved_docs}


def generate(state: State):
    docs_content = "\n\n".join(doc.page_content for doc in state["context"])
    messages = rag_prompt.invoke({"question": state["question"], "context": docs_content})
    response = llm.invoke(messages)
    return {"answer": response.content}


graph_builder = StateGraph(State).add_sequence([retrieve, generate])
graph_builder.add_edge(START, "retrieve")
graph = graph_builder.compile()


response = graph.invoke({"question": "Query"})
print(response["answer"])