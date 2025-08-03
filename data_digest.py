from langchain_core.documents import Document
from langchain_community.document_loaders import PyMuPDFLoader
from docx import Document as DocxDocument
from PIL import Image
import pytesseract
import pandas as pd
import sqlite3
import os


def load_pdf(file_path):
    loader = PyMuPDFLoader(file_path)
    return loader.load()

def load_docx(file_path):
    doc = DocxDocument(file_path)
    full_text = "\n".join([para.text for para in doc.paragraphs])
    return [Document(page_content=full_text, metadata={"source": file_path})]

def load_txt(file_path):
    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read()
    return [Document(page_content=text, metadata={"source": file_path})]

def load_image(file_path):
    text = pytesseract.image_to_string(Image.open(file_path))
    return [Document(page_content=text, metadata={"source": file_path})]

def load_csv(file_path):
    df = pd.read_csv(file_path)
    text = df.to_csv(index=False)
    return [Document(page_content=text, metadata={"source": file_path})]

def load_sqlite(file_path, query="SELECT name FROM sqlite_master WHERE type='table';"):
    conn = sqlite3.connect(file_path)
    cursor = conn.cursor()
    documents = []
    cursor.execute(query)
    tables = [row[0] for row in cursor.fetchall()]
    for table in tables:
        df = pd.read_sql_query(f"SELECT * FROM {table}", conn)
        text = df.to_csv(index=False)
        documents.append(Document(page_content=text, metadata={"source": f"{file_path}:{table}"}))
    conn.close()
    return documents

def load_sql_file(file_path):
    with open(file_path, "r", encoding="utf-8") as f:
        sql_text = f.read()
    return [Document(page_content=sql_text, metadata={"source": file_path})]

def load_xlsx(file_path):
    df = pd.read_excel(file_path)
    text = df.to_csv(index=False)
    return [Document(page_content=text, metadata={"source": file_path})]

def load_json(file_path):
    with open(file_path, "r", encoding="utf-8") as f:
        data = f.read()
    return [Document(page_content=data, metadata={"source": file_path})]

def load_md(file_path):
    with open(file_path, "r", encoding="utf-8") as f:
        markdown_text = f.read()
    return [Document(page_content=markdown_text, metadata={"source": file_path})]


def load_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return load_pdf(file_path)
    elif ext == ".docx":
        return load_docx(file_path)
    elif ext == ".txt":
        return load_txt(file_path)
    elif ext in [".jpg", ".jpeg", ".png"]:
        return load_image(file_path)
    elif ext == ".csv":
        return load_csv(file_path)
    elif ext in [".db", ".sqlite"]:
        return load_sqlite(file_path)
    elif ext == ".sql":
        return load_sql_file(file_path)
    elif ext == ".xlsx":
        return load_xlsx(file_path)
    elif ext == ".json":
        return load_json(file_path)
    elif ext == ".md":
        return load_md(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")